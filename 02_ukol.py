from docx import Document

def create_protokol(doc_path, argument1, argument2, argument3):
    
    doc = Document()


    title = doc.add_heading('Výsledný protokol genetického vyšetření', level=1)
    title.alignment = 1

    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Jméno a příjmení:'
    table.cell(0, 1).text = argument1
    table.cell(1, 0).text = 'Rodné číslo:'
    table.cell(1, 1).text = argument2
    table.cell(2, 0).text = 'Datum odběru:'
    table.cell(2, 1).text = argument3

    
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = 1 

    doc.save(doc_path)


create_protokol('Protokol.docx', 'argument1', 'argument2','argument3')